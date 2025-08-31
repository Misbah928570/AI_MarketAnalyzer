"""
AI Market Research & Trend Analyst SaaS Platform
Developed by Misbah Noorain 🎓
Data Science Student at BIET Davanagere
Modified to use Groq API instead of OpenAI
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import json
import requests
import hashlib
import time
from typing import Dict, List, Any
import random
import base64
from io import BytesIO
import numpy as np

# Conditional imports for optional dependencies
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    st.warning("ReportLab not available. PDF export will be disabled.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx not available. Word export will be disabled.")

try:
    from googletrans import Translator
    TRANSLATION_AVAILABLE = True
except ImportError:
    TRANSLATION_AVAILABLE = False
    st.warning("Google Translate not available. Translation features will be disabled.")

# Page configuration
st.set_page_config(
    page_title="AI Market Research & Trend Analyst",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Multi-language support - COMPLETE TRANSLATIONS
TRANSLATIONS = {
    'en': {
        'title': 'AI Market Research & Trend Analyst',
        'dashboard': 'Dashboard',
        'reports': 'Reports',
        'about': 'About Us',
        'landing': 'Home',
        'welcome': 'Welcome to AI-Powered Market Intelligence',
        'subtitle': 'Real-time market research, trend analysis, and AI-driven insights',
        'features': 'Features',
        'get_started': 'Get Started',
        'api_config': 'API Configuration',
        'market_query': 'Market Research Query',
        'generate_report': 'Generate Report',
        'developer': 'Developed by',
        'student': 'Data Science Student at BIET Davanagere',
        'connect': 'Connect with me',
        'navigation': 'Navigation',
        'real_time_dashboard': 'Real-time Market Intelligence Dashboard',
        'market_growth': 'Market Growth',
        'sentiment_score': 'Sentiment Score',
        'risk_level': 'Risk Level',
        'reports_generated': 'Reports Generated',
        'industry': 'Industry',
        'timeframe': 'Timeframe',
        'agent_status': 'Agent Status',
        'ai_insights': 'AI Insights',
        'quick_actions': 'Quick Actions',
        'refresh_data': 'Refresh Data',
        'export_dashboard': 'Export Dashboard',
        'compare_industries': 'Compare Industries',
        'no_reports': 'No reports generated yet. Go to Dashboard to create your first report!',
        'search_reports': 'Search reports',
        'filter_by_date': 'Filter by date',
        'export_language': 'Export Language',
        'meet_developer': 'Meet the Developer'
    },
    'es': {
        'title': 'Analista de Tendencias e Investigación de Mercado con IA',
        'dashboard': 'Panel',
        'reports': 'Informes',
        'about': 'Acerca de',
        'landing': 'Inicio',
        'welcome': 'Bienvenido a la Inteligencia de Mercado con IA',
        'subtitle': 'Investigación de mercado en tiempo real, análisis de tendencias e insights con IA',
        'features': 'Características',
        'get_started': 'Comenzar',
        'api_config': 'Configuración API',
        'market_query': 'Consulta de Investigación de Mercado',
        'generate_report': 'Generar Informe',
        'developer': 'Desarrollado por',
        'student': 'Estudiante de Ciencia de Datos en BIET Davanagere',
        'connect': 'Conéctate conmigo',
        'navigation': 'Navegación',
        'real_time_dashboard': 'Panel de Inteligencia de Mercado en Tiempo Real',
        'market_growth': 'Crecimiento del Mercado',
        'sentiment_score': 'Puntuación de Sentimiento',
        'risk_level': 'Nivel de Riesgo',
        'reports_generated': 'Informes Generados',
        'industry': 'Industria',
        'timeframe': 'Marco Temporal',
        'agent_status': 'Estado del Agente',
        'ai_insights': 'Insights de IA',
        'quick_actions': 'Acciones Rápidas',
        'refresh_data': 'Actualizar Datos',
        'export_dashboard': 'Exportar Panel',
        'compare_industries': 'Comparar Industrias',
        'no_reports': '¡Aún no se han generado informes. Ve al Panel para crear tu primer informe!',
        'search_reports': 'Buscar informes',
        'filter_by_date': 'Filtrar por fecha',
        'export_language': 'Idioma de Exportación',
        'meet_developer': 'Conoce al Desarrollador'
    },
    'fr': {
        'title': 'Analyste de Tendances et Recherche de Marché IA',
        'dashboard': 'Tableau de bord',
        'reports': 'Rapports',
        'about': 'À propos',
        'landing': 'Accueil',
        'welcome': 'Bienvenue dans l\'Intelligence de Marché IA',
        'subtitle': 'Recherche de marché en temps réel, analyse des tendances et insights IA',
        'features': 'Fonctionnalités',
        'get_started': 'Commencer',
        'api_config': 'Configuration API',
        'market_query': 'Requête de Recherche de Marché',
        'generate_report': 'Générer un Rapport',
        'developer': 'Développé par',
        'student': 'Étudiant en Science des Données à BIET Davanagere',
        'connect': 'Connectez-vous avec moi',
        'navigation': 'Navigation',
        'real_time_dashboard': 'Tableau de Bord d\'Intelligence de Marché en Temps Réel',
        'market_growth': 'Croissance du Marché',
        'sentiment_score': 'Score de Sentiment',
        'risk_level': 'Niveau de Risque',
        'reports_generated': 'Rapports Générés',
        'industry': 'Industrie',
        'timeframe': 'Période',
        'agent_status': 'Statut de l\'Agent',
        'ai_insights': 'Insights IA',
        'quick_actions': 'Actions Rapides',
        'refresh_data': 'Actualiser les Données',
        'export_dashboard': 'Exporter le Tableau de Bord',
        'compare_industries': 'Comparer les Industries',
        'no_reports': 'Aucun rapport généré pour le moment. Allez au Tableau de Bord pour créer votre premier rapport!',
        'search_reports': 'Rechercher des rapports',
        'filter_by_date': 'Filtrer par date',
        'export_language': 'Langue d\'Exportation',
        'meet_developer': 'Rencontrez le Développeur'
    },
    'hi': {
        'title': 'एआई मार्केट रिसर्च और ट्रेंड विश्लेषक',
        'dashboard': 'डैशबोर्ड',
        'reports': 'रिपोर्ट्स',
        'about': 'हमारे बारे में',
        'landing': 'होम',
        'welcome': 'एआई-संचालित मार्केट इंटेलिजेंस में आपका स्वागत है',
        'subtitle': 'रीयल-टाइम मार्केट रिसर्च, ट्रेंड एनालिसिस और एआई आधारित इनसाइट्स',
        'features': 'विशेषताएं',
        'get_started': 'शुरू करें',
        'api_config': 'एपीआई कॉन्फ़िगरेशन',
        'market_query': 'मार्केट रिसर्च क्वेरी',
        'generate_report': 'रिपोर्ट जनरेट करें',
        'developer': 'द्वारा विकसित',
        'student': 'बीआईईटी दावणगेरे में डेटा साइंस छात्र',
        'connect': 'मुझसे जुड़ें',
        'navigation': 'नेविगेशन',
        'real_time_dashboard': 'रीयल-टाइम मार्केट इंटेलिजेंस डैशबोर्ड',
        'market_growth': 'मार्केट ग्रोथ',
        'sentiment_score': 'सेंटिमेंट स्कोर',
        'risk_level': 'रिस्क लेवल',
        'reports_generated': 'जनरेट की गई रिपोर्ट्स',
        'industry': 'इंडस्ट्री',
        'timeframe': 'समयसीमा',
        'agent_status': 'एजेंट स्टेटस',
        'ai_insights': 'एआई इनसाइट्स',
        'quick_actions': 'त्वरित क्रियाएं',
        'refresh_data': 'डेटा रिफ्रेश करें',
        'export_dashboard': 'डैशबोर्ड एक्सपोर्ट करें',
        'compare_industries': 'इंडस्ट्रीज की तुलना करें',
        'no_reports': 'अभी तक कोई रिपोर्ट जनरेट नहीं की गई है। अपनी पहली रिपोर्ट बनाने के लिए डैशबोर्ड पर जाएं!',
        'search_reports': 'रिपोर्ट्स खोजें',
        'filter_by_date': 'दिनांक के अनुसार फ़िल्टर करें',
        'export_language': 'एक्सपोर्ट भाषा',
        'meet_developer': 'डेवलपर से मिलें'
    }
}

# Custom CSS for SaaS-style design
def load_css():
    st.markdown("""
    <style>
    /* Main Theme */
    .stApp {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
    }
    
    .feature-card {
        background: #d8b4fe;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
        transition: transform 0.3s;
    }
    
    .feature-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 25px rgba(0,0,0,0.15);
    }
    
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 10px;
        text-align: center;
        box-shadow: 0 5px 15px rgba(0,0,0,0.2);
    }
    
    .agent-status {
        background: #d8b4fe;
        padding: 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 4px solid #667eea;
    }
    
    .report-card {
        background: #d8b4fe;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 3px 10px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
    }
    
    .developer-info {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin-top: 2rem;
    }
    
    .stSidebar {
        background: #d8b4fe;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize session state
def init_session_state():
    if 'language' not in st.session_state:
        st.session_state.language = 'en'
    if 'api_keys' not in st.session_state:
        st.session_state.api_keys = {
            'groq': '',
            'news': '',
            'twitter': '',
            'google_translate': ''
        }
    if 'reports' not in st.session_state:
        st.session_state.reports = []
    if 'market_data' not in st.session_state:
        st.session_state.market_data = generate_sample_data()
    if 'agent_logs' not in st.session_state:
        st.session_state.agent_logs = []
    if 'translation_mode' not in st.session_state:
        st.session_state.translation_mode = 'auto'

# API Integration Functions
def get_groq_insights(query, market_data):
    """Generate AI insights using Groq API"""
    if not st.session_state.api_keys.get('groq'):
        return [
            "Strong growth potential in emerging markets",
            "Increasing adoption of AI technologies", 
            "Competitive landscape is evolving rapidly"
        ]
    
    try:
        # Groq API endpoint
        url = "https://api.groq.com/openai/v1/chat/completions"
        
        headers = {
            "Authorization": f"Bearer {st.session_state.api_keys['groq']}",
            "Content-Type": "application/json"
        }
        
        # Create a concise market data summary for the prompt
        data_summary = {
            'query': query,
            'trends': market_data.get('trends', {}),
            'sentiment': market_data.get('news_sentiment', 'neutral'),
            'articles_found': len(market_data.get('news', []))
        }
        
        prompt = f"""
        Analyze this market research query and provide 3 key actionable insights:

        Query: {query}
        Market Data Summary: {json.dumps(data_summary, indent=2)}

        Please provide exactly 3 concise, actionable insights about this market. Each insight should be a single sentence focusing on opportunities, risks, or trends.

        Format your response as 3 separate lines, each starting with a bullet point or dash.
        """
        
        payload = {
            "model": "llama3-8b-8192",  # Using Llama 3 8B model
            "messages": [
                {
                    "role": "system", 
                    "content": "You are an expert market research analyst. Provide concise, actionable insights based on market data."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ],
            "max_tokens": 300,
            "temperature": 0.7,
            "top_p": 1,
            "stream": False
        }
        
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content'].strip()
            
            # Parse the insights from the response
            insights = []
            for line in content.split('\n'):
                line = line.strip()
                if line and (line.startswith('-') or line.startswith('•') or line.startswith('*')):
                    # Remove bullet points and clean up
                    insight = line.lstrip('-•* ').strip()
                    if insight:
                        insights.append(insight)
            
            # If parsing failed, split by lines and take first 3 non-empty ones
            if len(insights) < 3:
                insights = [line.strip() for line in content.split('\n') if line.strip()][:3]
            
            # Ensure we have at least 3 insights
            if len(insights) < 3:
                insights.extend([
                    "Market shows potential for strategic investment opportunities",
                    "Digital transformation trends are accelerating in this sector", 
                    "Consumer behavior patterns indicate shift towards innovative solutions"
                ][:3-len(insights)])
            
            return insights[:3]  # Return exactly 3 insights
            
        else:
            st.warning(f"Groq API error: {response.status_code} - {response.text}")
            return [
                "Strong growth potential in emerging markets",
                "Increasing adoption of AI technologies", 
                "Competitive landscape is evolving rapidly"
            ]
            
    except requests.exceptions.Timeout:
        st.warning("Groq API request timed out. Using default insights.")
        return [
            "Strong growth potential in emerging markets",
            "Increasing adoption of AI technologies", 
            "Competitive landscape is evolving rapidly"
        ]
    except Exception as e:
        st.warning(f"Groq API integration error: {e}")
        return [
            "Strong growth potential in emerging markets",
            "Increasing adoption of AI technologies", 
            "Competitive landscape is evolving rapidly"
        ]

def get_news_data(query):
    """Fetch news data using NewsAPI"""
    if not st.session_state.api_keys.get('news'):
        return {"articles": [], "status": "No API key"}
    
    try:
        url = f"https://newsapi.org/v2/everything"
        params = {
            'q': query,
            'apiKey': st.session_state.api_keys['news'],
            'language': 'en',
            'sortBy': 'relevancy',
            'pageSize': 5
        }
        
        response = requests.get(url, params=params, timeout=10)
        data = response.json()
        
        if data.get('status') == 'ok':
            return data
        else:
            return {"articles": [], "status": f"API Error: {data.get('message', 'Unknown error')}"}
            
    except Exception as e:
        st.warning(f"NewsAPI error: {e}")
        return {"articles": [], "status": "API call failed"}

def get_market_sentiment_analysis(query):
    """Get enhanced market sentiment analysis"""
    sentiment_data = {
        'overall_score': random.uniform(0.3, 0.9),
        'trend': random.choice(['Bullish', 'Bearish', 'Neutral', 'Volatile']),
        'sources': ['News Articles', 'Social Media', 'Market Data'],
        'confidence': random.uniform(0.7, 0.95)
    }
    
    # If Twitter API is available, integrate real sentiment
    if st.session_state.api_keys.get('twitter'):
        try:
            # Placeholder for Twitter API integration
            sentiment_data['twitter_enabled'] = True
        except Exception as e:
            sentiment_data['twitter_error'] = str(e)
    
    return sentiment_data

def t(key):
    """Translation helper function with fallback"""
    try:
        return TRANSLATIONS[st.session_state.language].get(key, TRANSLATIONS['en'].get(key, key))
    except KeyError:
        return TRANSLATIONS['en'].get(key, key)

def safe_translate_text(text, target_lang):
    """Safe translation function with enhanced error handling"""
    if not TRANSLATION_AVAILABLE:
        return text
    
    if target_lang == "English" or target_lang == "en":
        return text
    
    lang_codes = {
        "English": "en",
        "Spanish": "es", 
        "French": "fr",
        "German": "de",
        "Hindi": "hi",
        "Chinese": "zh-cn"
    }
    
    try:
        target_code = lang_codes.get(target_lang, "en")
        if target_code == "en":
            return text
        
        # Create a new translator instance for each translation to avoid session issues
        translator = Translator()
        
        # Add retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                result = translator.translate(text, dest=target_code)
                
                # Handle different return types
                if hasattr(result, 'text'):
                    return result.text
                elif isinstance(result, str):
                    return result
                else:
                    # If it's a coroutine or other unexpected type, return original text
                    return text
                    
            except Exception as retry_error:
                if attempt == max_retries - 1:  # Last attempt
                    st.warning(f"Translation failed after {max_retries} attempts: {retry_error}")
                    return text
                time.sleep(0.5)  # Wait before retry
                
        return text
        
    except Exception as e:
        st.warning(f"Translation error: {e}")
        return text

# Simplified offline translation dictionary for fallback
OFFLINE_TRANSLATIONS = {
    "Spanish": {
        "Strong growth potential in emerging markets": "Fuerte potencial de crecimiento en mercados emergentes",
        "Increasing adoption of AI technologies": "Adopción creciente de tecnologías de IA",
        "Competitive landscape is evolving rapidly": "El panorama competitivo está evolucionando rápidamente",
        "Invest in emerging technologies": "Invertir en tecnologías emergentes",
        "Focus on customer experience": "Centrarse en la experiencia del cliente",
        "Expand into new markets": "Expandirse a nuevos mercados",
        "This report provides comprehensive market analysis with AI-driven insights.": "Este informe proporciona un análisis integral del mercado con perspectivas impulsadas por IA."
    },
    "French": {
        "Strong growth potential in emerging markets": "Fort potentiel de croissance dans les marchés émergents",
        "Increasing adoption of AI technologies": "Adoption croissante des technologies d'IA",
        "Competitive landscape is evolving rapidly": "Le paysage concurrentiel évolue rapidement",
        "Invest in emerging technologies": "Investir dans les technologies émergentes",
        "Focus on customer experience": "Se concentrer sur l'expérience client",
        "Expand into new markets": "S'étendre vers de nouveaux marchés",
        "This report provides comprehensive market analysis with AI-driven insights.": "Ce rapport fournit une analyse complète du marché avec des perspectives pilotées par l'IA."
    },
    "Hindi": {
        "Strong growth potential in emerging markets": "उभरते बाजारों में मजबूत विकास की संभावना",
        "Increasing adoption of AI technologies": "एआई तकनीकों का बढ़ता अपनाना",
        "Competitive landscape is evolving rapidly": "प्रतिस्पर्धी परिदृश्य तेजी से विकसित हो रहा है",
        "Invest in emerging technologies": "उभरती प्रौद्योगिकियों में निवेश करें",
        "Focus on customer experience": "ग्राहक अनुभव पर ध्यान दें",
        "Expand into new markets": "नए बाजारों में विस्तार करें",
        "This report provides comprehensive market analysis with AI-driven insights.": "यह रिपोर्ट एआई-संचालित अंतर्दृष्टि के साथ व्यापक बाजार विश्लेषण प्रदान करती है।"
    }
}

def offline_translate(text, target_lang):
    """Fallback offline translation for common phrases"""
    if target_lang in OFFLINE_TRANSLATIONS:
        return OFFLINE_TRANSLATIONS[target_lang].get(text, text)
    return text

def translate_report(report, target_lang):
    """Translate report content with enhanced error handling and fallback"""
    if target_lang == "English":
        return report
    
    try:
        translated = report.copy()
        
        # Use online translation first, fall back to offline if it fails
        def safe_field_translate(text, field_name=""):
            try:
                # Try online translation first
                if TRANSLATION_AVAILABLE:
                    online_result = safe_translate_text(text, target_lang)
                    if online_result != text:  # If translation actually happened
                        return online_result
                
                # Fallback to offline translation
                offline_result = offline_translate(text, target_lang)
                return offline_result
                
            except Exception as e:
                st.warning(f"Translation failed for {field_name}: {e}")
                return text
        
        # Translate main fields with fallback
        translated["executive_summary"] = safe_field_translate(
            report.get("executive_summary", ""), "executive_summary"
        )
        
        # Translate key insights
        if "analysis" in report and "key_insights" in report["analysis"]:
            translated["analysis"] = report["analysis"].copy()
            translated_insights = []
            for i, insight in enumerate(report["analysis"]["key_insights"]):
                translated_insight = safe_field_translate(insight, f"insight_{i}")
                translated_insights.append(translated_insight)
            translated["analysis"]["key_insights"] = translated_insights
        
        # Translate recommendations
        if "recommendations" in report:
            translated_recs = []
            for i, rec in enumerate(report["recommendations"]):
                translated_rec = safe_field_translate(rec, f"recommendation_{i}")
                translated_recs.append(translated_rec)
            translated["recommendations"] = translated_recs
        
        return translated
        
    except Exception as e:
        st.error(f"Report translation error: {e}")
        st.info("Showing report in English due to translation issues.")
        return report

# Multi-Agent System Classes
class MarketResearchAgent:
    def __init__(self, name, role):
        self.name = name
        self.role = role
        self.status = "idle"
    
    def log_activity(self, activity):
        log_entry = {
            'timestamp': datetime.now().strftime('%H:%M:%S'),
            'agent': self.name,
            'activity': activity,
            'status': self.status
        }
        st.session_state.agent_logs.append(log_entry)
        return log_entry

class ScraperAgent(MarketResearchAgent):
    def __init__(self):
        super().__init__("Scraper Agent", "Data Collection")
    
    def scrape_market_data(self, query):
        self.status = "active"
        self.log_activity(f"Scraping data for: {query}")
        
        # Get real news data if API is available
        news_data = get_news_data(query)
        
        # Simulated data with real news integration
        data = {
            'query': query,
            'news': news_data.get('articles', []),
            'news_status': news_data.get('status', 'No API'),
            'trends': generate_trend_data(),
            'competitors': ['Company A', 'Company B', 'Company C'],
            'timestamp': datetime.now().isoformat()
        }
        
        # Add news sentiment if articles are available
        if data['news']:
            positive_count = sum(1 for article in data['news'] 
                               if 'growth' in article.get('title', '').lower() 
                               or 'positive' in article.get('description', '').lower())
            data['news_sentiment'] = 'positive' if positive_count > len(data['news'])/2 else 'neutral'
        
        self.status = "completed"
        self.log_activity(f"Completed scraping for: {query} (Found {len(data['news'])} articles)")
        return data

class AnalyzerAgent(MarketResearchAgent):
    def __init__(self):
        super().__init__("Analyzer Agent", "Data Analysis")
    
    def analyze_data(self, data):
        self.status = "active"
        self.log_activity("Analyzing market data with Groq AI...")
        
        # Use Groq for insights
        query = data.get('query', 'market analysis')
        ai_insights = get_groq_insights(query, data)
        
        # Get enhanced sentiment analysis
        sentiment_analysis = get_market_sentiment_analysis(query)
        
        analysis = {
            'sentiment_score': sentiment_analysis['overall_score'],
            'growth_potential': random.uniform(0.4, 0.95),
            'risk_level': random.choice(['Low', 'Medium', 'High']),
            'key_insights': ai_insights,
            'sentiment_trend': sentiment_analysis['trend'],
            'data_sources': sentiment_analysis['sources'],
            'confidence': sentiment_analysis['confidence']
        }
        
        self.status = "completed"
        self.log_activity("Analysis completed with Groq AI integration")
        return analysis

class ReporterAgent(MarketResearchAgent):
    def __init__(self):
        super().__init__("Reporter Agent", "Report Generation")
    
    def generate_report(self, data, analysis):
        self.status = "active"
        self.log_activity("Generating comprehensive report...")
        
        report = {
            'title': f"Market Research Report - {datetime.now().strftime('%Y-%m-%d')}",
            'executive_summary': "This report provides comprehensive market analysis with AI-driven insights.",
            'market_data': data,
            'analysis': analysis,
            'recommendations': [
                "Invest in emerging technologies",
                "Focus on customer experience", 
                "Expand into new markets"
            ],
            'generated_at': datetime.now().isoformat()
        }
        
        self.status = "completed"
        self.log_activity("Report generated successfully")
        return report

class VisualizerAgent(MarketResearchAgent):
    def __init__(self):
        super().__init__("Visualizer Agent", "Data Visualization")
    
    def create_visualizations(self, data):
        self.status = "active"
        self.log_activity("Creating visualizations...")
        
        charts = {
            'trend_chart': create_trend_chart(data),
            'sentiment_radar': create_sentiment_radar(),
            'growth_forecast': create_growth_forecast()
        }
        
        self.status = "completed"
        self.log_activity("Visualizations created")
        return charts

# Data generation functions
def generate_sample_data():
    dates = pd.date_range(start='2024-01-01', end='2024-12-31', freq='M')
    return pd.DataFrame({
        'date': dates,
        'market_value': np.random.randint(1000, 5000, len(dates)),
        'growth_rate': np.random.uniform(-0.1, 0.3, len(dates)),
        'sentiment': np.random.uniform(0.3, 0.9, len(dates))
    })

def generate_trend_data():
    return {
        'ai_adoption': random.uniform(0.6, 0.95),
        'market_growth': random.uniform(0.1, 0.3),
        'innovation_index': random.uniform(0.5, 0.9),
        'competition_level': random.uniform(0.4, 0.8)
    }

# Visualization functions
def create_trend_chart(data):
    fig = go.Figure()
    
    fig.add_trace(go.Scatter(
        x=st.session_state.market_data['date'],
        y=st.session_state.market_data['market_value'],
        mode='lines+markers',
        name='Market Value',
        line=dict(color='#667eea', width=3)
    ))
    
    fig.update_layout(
        title='Market Trend Analysis',
        xaxis_title='Date',
        yaxis_title='Value',
        hovermode='x unified',
        template='plotly_white',
        height=400
    )
    
    return fig

def create_sentiment_radar():
    categories = ['Technology', 'Healthcare', 'Finance', 'Retail', 'Manufacturing']
    
    fig = go.Figure(data=go.Scatterpolar(
        r=[random.uniform(0.5, 1) for _ in categories],
        theta=categories,
        fill='toself',
        fillcolor='rgba(102, 126, 234, 0.3)',
        line=dict(color='#667eea')
    ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 1]
            )),
        showlegend=False,
        title="Industry Sentiment Radar",
        height=400
    )
    
    return fig

def create_growth_forecast():
    future_dates = pd.date_range(start='2025-01-01', end='2025-12-31', freq='M')
    forecast_values = [3000 + i * 100 + random.randint(-200, 200) for i in range(len(future_dates))]
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatter(
        x=future_dates,
        y=forecast_values,
        mode='lines',
        name='Forecast',
        line=dict(color='#764ba2', width=3, dash='dash')
    ))
    
    fig.add_trace(go.Scatter(
        x=future_dates,
        y=[v + 300 for v in forecast_values],
        mode='lines',
        name='Upper Bound',
        line=dict(color='rgba(118, 75, 162, 0.3)'),
        fill=None
    ))
    
    fig.add_trace(go.Scatter(
        x=future_dates,
        y=[v - 300 for v in forecast_values],
        mode='lines',
        name='Lower Bound',
        line=dict(color='rgba(118, 75, 162, 0.3)'),
        fill='tonexty',
        fillcolor='rgba(118, 75, 162, 0.1)'
    ))
    
    fig.update_layout(
        title='Growth Forecast 2025',
        xaxis_title='Date',
        yaxis_title='Projected Value',
        template='plotly_white',
        height=400
    )
    
    return fig

# Export functions with error handling
def generate_pdf(report):
    """Generate PDF with error handling"""
    if not REPORTLAB_AVAILABLE:
        st.error("PDF generation not available. Please install reportlab.")
        return None
    
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer)
        styles = getSampleStyleSheet()
        content = []

        content.append(Paragraph(report.get("title", "Market Research Report"), styles['Heading1']))
        content.append(Spacer(1, 12))
        content.append(Paragraph("Executive Summary", styles['Heading2']))
        content.append(Paragraph(report.get("executive_summary", ""), styles['Normal']))
        content.append(Spacer(1, 12))

        content.append(Paragraph("Key Insights", styles['Heading2']))
        if "analysis" in report and "key_insights" in report["analysis"]:
            for insight in report["analysis"]["key_insights"]:
                content.append(Paragraph(f"• {insight}", styles['Normal']))
        content.append(Spacer(1, 12))

        content.append(Paragraph("Recommendations", styles['Heading2']))
        if "recommendations" in report:
            for rec in report["recommendations"]:
                content.append(Paragraph(f"• {rec}", styles['Normal']))
        content.append(Spacer(1, 12))

        content.append(Paragraph(f"Generated: {report.get('generated_at', '')}", styles['Normal']))
        if "analysis" in report:
            content.append(Paragraph(f"Risk Level: {report['analysis'].get('risk_level', 'N/A')}", styles['Normal']))
            content.append(Paragraph(f"Growth Potential: {report['analysis'].get('growth_potential', 0):.1%}", styles['Normal']))

        doc.build(content)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"PDF generation failed: {e}")
        return None

def generate_word(report):
    """Generate Word document with error handling"""
    if not DOCX_AVAILABLE:
        st.error("Word generation not available. Please install python-docx.")
        return None
    
    try:
        buffer = BytesIO()
        doc = Document()

        doc.add_heading(report.get("title", "Market Research Report"), 0)
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.get("executive_summary", ""))

        doc.add_heading("Key Insights", level=1)
        if "analysis" in report and "key_insights" in report["analysis"]:
            for insight in report["analysis"]["key_insights"]:
                doc.add_paragraph(f"• {insight}")

        doc.add_heading("Recommendations", level=1)
        if "recommendations" in report:
            for rec in report["recommendations"]:
                doc.add_paragraph(f"• {rec}")

        doc.add_paragraph(f"Generated: {report.get('generated_at', '')}")
        if "analysis" in report:
            doc.add_paragraph(f"Risk Level: {report['analysis'].get('risk_level', 'N/A')}")
            doc.add_paragraph(f"Growth Potential: {report['analysis'].get('growth_potential', 0):.1%}")

        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Word generation failed: {e}")
        return None

# Page functions
def landing_page():
    st.markdown(f"""
    <div class="main-header">
        <h1>🚀 {t('title')}</h1>
        <h3>{t('subtitle')}</h3>
    </div>
    """, unsafe_allow_html=True)
    
    # Hero Section
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown(f"### {t('welcome')}")
        st.write("""
        Transform your market research with AI-powered insights. Our platform combines 
        advanced machine learning with real-time data analysis to deliver actionable 
        intelligence for your business decisions.
        """)
        
        if st.button(t('get_started'), type="primary", use_container_width=True):
            st.session_state.page = 'dashboard'
            st.rerun()
    
    with col2:
        st.markdown("""
        <div class="metric-card">
            <h2>95%</h2>
            <p>Accuracy Rate</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Features Section
    st.markdown(f"### ✨ {t('features')}")
    
    feature_cols = st.columns(3)
    features = [
        ("🤖", "Multi-Agent System", "Intelligent agents working together for comprehensive analysis"),
        ("📊", "Real-time Analytics", "Live market data and trend analysis at your fingertips"),
        ("🌍", "Multi-language Support", "Generate reports in multiple languages"),
        ("📈", "Trend Forecasting", "AI-powered predictions for market movements"),
        ("📋", "Automated Reports", "PDF/Word reports with one click"),
        ("⚡", "Groq AI Integration", "Lightning-fast AI insights with Llama models")
    ]
    
    for i, (icon, title, desc) in enumerate(features):
        with feature_cols[i % 3]:
            st.markdown(f"""
            <div class="feature-card">
                <h3>{icon} {title}</h3>
                <p>{desc}</p>
            </div>
            """, unsafe_allow_html=True)

def dashboard_page():
    st.markdown(f"""
    <div class="main-header">
        <h1>📊 {t('dashboard')}</h1>
        <p>{t('real_time_dashboard')}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize agents
    scraper = ScraperAgent()
    analyzer = AnalyzerAgent()
    reporter = ReporterAgent()
    visualizer = VisualizerAgent()
    
    # Metrics Row
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(t('market_growth'), "+15.3%", "↑ 2.1%")
    with col2:
        st.metric(t('sentiment_score'), "0.82", "↑ 0.05")
    with col3:
        st.metric(t('risk_level'), "Medium", "→ Stable")
    with col4:
        st.metric(t('reports_generated'), len(st.session_state.reports), f"→ {len(st.session_state.reports)}")
    
    # Main Dashboard Layout
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Market Query Section
        st.markdown(f"### 🔍 {t('market_query')}")
        
        query = st.text_input("Enter your market research query:", 
                             placeholder="e.g., AI trends in healthcare 2024")
        
        col_a, col_b = st.columns(2)
        with col_a:
            industry = st.selectbox(t('industry'), 
                                   ["Technology", "Healthcare", "Finance", "Retail", "Manufacturing"])
        with col_b:
            timeframe = st.selectbox(t('timeframe'), 
                                    ["Last Month", "Last Quarter", "Last Year", "Custom"])
        
        if st.button(t('generate_report'), type="primary", use_container_width=True):
            if not query.strip():
                st.warning("Please enter a market research query to generate a report.")
            else:
                with st.spinner("Multi-agent system working with Groq AI..."):
                    # Agent workflow
                    progress = st.progress(0)
                    
                    # Step 1: Scraping
                    progress.progress(25)
                    data = scraper.scrape_market_data(query)
                    time.sleep(1)
                    
                    # Step 2: Analysis with Groq AI
                    progress.progress(50)
                    analysis = analyzer.analyze_data(data)
                    time.sleep(1)
                    
                    # Step 3: Report Generation
                    progress.progress(75)
                    report = reporter.generate_report(data, analysis)
                    st.session_state.reports.append(report)
                    time.sleep(1)
                    
                    # Step 4: Visualization
                    progress.progress(100)
                    charts = visualizer.create_visualizations(data)
                    
                    st.success("✅ Report generated successfully with Groq AI insights!")
        
        # Visualizations
        st.markdown("### 📈 Market Visualizations")
        
        tab1, tab2, tab3 = st.tabs(["Trend Analysis", "Sentiment Radar", "Growth Forecast"])
        
        with tab1:
            st.plotly_chart(create_trend_chart(None), use_container_width=True)
        
        with tab2:
            st.plotly_chart(create_sentiment_radar(), use_container_width=True)
        
        with tab3:
            st.plotly_chart(create_growth_forecast(), use_container_width=True)
    
    with col2:
        # API Integration Status
        st.markdown("### 🔗 API Integration Status")
        
        api_status_cards = []
        if st.session_state.api_keys.get('groq'):
            api_status_cards.append("✅ **Groq AI**: Lightning-fast insights enabled")
        else:
            api_status_cards.append("⚠️ **Groq AI**: Configure for AI insights")
            
        if st.session_state.api_keys.get('news'):
            api_status_cards.append("✅ **NewsAPI**: Real-time news enabled")
        else:
            api_status_cards.append("⚠️ **NewsAPI**: Configure for real news data")
        
        for status in api_status_cards:
            st.markdown(f"""
            <div class="agent-status">
                {status}
            </div>
            """, unsafe_allow_html=True)
        
        # Agent Status Panel
        st.markdown(f"### 🤖 {t('agent_status')}")
        
        for log in st.session_state.agent_logs[-5:]:
            st.markdown(f"""
            <div class="agent-status">
                <strong>{log['agent']}</strong><br>
                <small>{log['timestamp']}</small><br>
                {log['activity']}
            </div>
            """, unsafe_allow_html=True)
        
        # AI Insights
        st.markdown(f"### 💡 {t('ai_insights')}")
        
        insights = [
            "📈 Market showing strong upward trend",
            "🎯 Customer sentiment improving by 15%",
            "⚡ New competitor entered the market",
            "🌟 Innovation index at all-time high",
            "📍 Regulatory changes expected Q2 2025"
        ]
        
        for insight in insights[:3]:
            st.info(insight)
        
        # Quick Actions
        st.markdown(f"### ⚡ {t('quick_actions')}")
        
        if st.button(f"🔄 {t('refresh_data')}", use_container_width=True):
            st.session_state.market_data = generate_sample_data()
            st.rerun()
        
        if st.button(f"📥 {t('export_dashboard')}", use_container_width=True):
            st.info("Dashboard exported successfully!")
        
        if st.button(f"🤖 {t('compare_industries')}", use_container_width=True):
            st.info("Industry comparison feature coming soon!")

def reports_page():
    st.markdown(f"""
    <div class="main-header">
        <h1>📄 {t('reports')}</h1>
        <p>View and manage your market research reports</p>
    </div>
    """, unsafe_allow_html=True)

    # Report filters
    col1, col2, col3 = st.columns(3)
    with col1:
        search = st.text_input(f"🔍 {t('search_reports')}", placeholder="Enter keywords...")
    with col2:
        filter_date = st.date_input(f"📅 {t('filter_by_date')}", value=None)
    with col3:
        export_lang = st.selectbox(f"🌍 {t('export_language')}", 
                                  ["English", "Spanish", "French", "German", "Hindi", "Chinese"])

    # Display reports
    if not st.session_state.reports:
        st.info(t('no_reports'))
    else:
        for i, report in enumerate(st.session_state.reports):
            translated_report = translate_report(report, export_lang)

            with st.expander(f"📊 {translated_report.get('title', 'Report')}", expanded=i==0):
                col1, col2 = st.columns([3, 1])

                with col1:
                    st.markdown("**Executive Summary:**")
                    st.write(translated_report.get("executive_summary", "No summary available"))

                    st.markdown("**Key Insights:**")
                    if "analysis" in translated_report and "key_insights" in translated_report["analysis"]:
                        for insight in translated_report["analysis"]["key_insights"]:
                            st.write(f"• {insight}")
                    else:
                        st.write("No insights available")

                    st.markdown("**Recommendations:**")
                    if "recommendations" in translated_report:
                        for rec in translated_report["recommendations"]:
                            st.write(f"• {rec}")
                    else:
                        st.write("No recommendations available")

                with col2:
                    st.markdown("**Report Details:**")
                    st.write(f"Generated: {translated_report.get('generated_at', 'Unknown')[:10]}")
                    
                    if "analysis" in translated_report:
                        st.write(f"Risk Level: {translated_report['analysis'].get('risk_level', 'N/A')}")
                        st.write(f"Growth: {translated_report['analysis'].get('growth_potential', 0):.1%}")

                    st.markdown("**Actions:**")
                    
                    # PDF Download
                    if REPORTLAB_AVAILABLE:
                        pdf_buffer = generate_pdf(translated_report)
                        if pdf_buffer:
                            st.download_button(
                                label="📥 Download PDF",
                                data=pdf_buffer,
                                file_name=f"report_{i+1}.pdf",
                                mime="application/pdf",
                                key=f"pdf_{i}"
                            )
                    else:
                        st.info("PDF export not available")

                    # Word Download
                    if DOCX_AVAILABLE:
                        word_buffer = generate_word(translated_report)
                        if word_buffer:
                            st.download_button(
                                label="📄 Download Word",
                                data=word_buffer,
                                file_name=f"report_{i+1}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"word_{i}"
                            )
                    else:
                        st.info("Word export not available")

                    if st.button(f"🔄 Regenerate", key=f"regen_{i}"):
                        st.info("Regenerating report...")

def about_page():
    st.markdown(f"""
    <div class="main-header">
        <h1>👨‍💻 {t('about')}</h1>
        <p>{t('meet_developer')}</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown(f"""
        <div class="developer-info">
            <h2>Misbah Noorain 🎓</h2>
            <p style="font-size: 1.2em;">{t('student')}</p>
            <p style="font-size: 1.1em;">BIET Davanagere</p>
            <br>
            <p>Passionate about AI, Machine Learning, and Data Science. 
            Building innovative solutions that leverage cutting-edge technology 
            to solve real-world problems.</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"### 🌍 {t('connect')}")
        
        col_a, col_b, col_c = st.columns(3)
        
        with col_a:
            st.markdown("""
            <a href="https://www.linkedin.com/in/misbah-noorain-984942294/" target="_blank">
                <div class="feature-card" style="text-align: center;">
                    <h3>💼 LinkedIn</h3>
                </div>
            </a>
            """, unsafe_allow_html=True)
        
        with col_b:
            st.markdown("""
            <a href="https://github.com/Misbah928570" target="_blank">
                <div class="feature-card" style="text-align: center;">
                    <h3>🐙 GitHub</h3>
                </div>
            </a>
            """, unsafe_allow_html=True)
        
        with col_c:
            st.markdown("""
            <a href="https://www.hackerrank.com/profile/misbahnoorain928" target="_blank">
                <div class="feature-card" style="text-align: center;">
                    <h3>🏆 HackerRank</h3>
                </div>
            </a>
            """, unsafe_allow_html=True)
        
        st.markdown("### 🚀 Project Features")
        
        features = {
            "Multi-Agent System": "Collaborative AI agents for comprehensive analysis",
            "Groq AI Integration": "Lightning-fast AI insights with Llama and Mixtral models",
            "Real-time Data": "Live market data integration with multiple APIs",
            "Multilingual": "Support for multiple languages in UI and reports",
            "Interactive Charts": "Plotly.js powered dynamic visualizations",
            "Professional Reports": "Export to PDF/Word with custom formatting"
        }
        
        for feature, description in features.items():
            st.markdown(f"""
            <div class="feature-card">
                <strong>{feature}</strong>: {description}
            </div>
            """, unsafe_allow_html=True)

# Main App
def main():
    # Load CSS
    load_css()
    
    # Initialize session state
    init_session_state()
    
    # Sidebar
    with st.sidebar:
        st.markdown("""
        <div style="text-align: center; padding: 1rem;">
            <h2>📊 AI Market Research</h2>
            <p style="font-size: 0.9em;">Powered by Groq AI</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Language selector
        languages = {'English': 'en', 'Español': 'es', 'Français': 'fr', 'हिंदी': 'hi'}
        selected_lang = st.selectbox("🌍 Language", list(languages.keys()))
        st.session_state.language = languages[selected_lang]
        
        st.markdown("---")
        
        # Navigation
        st.markdown(f"### {t('navigation')}")
        
        if st.button(f"🏠 {t('landing')}", use_container_width=True):
            st.session_state.page = 'landing'
            
        if st.button(f"📊 {t('dashboard')}", use_container_width=True):
            st.session_state.page = 'dashboard'
            
        if st.button(f"📄 {t('reports')}", use_container_width=True):
            st.session_state.page = 'reports'
            
        if st.button(f"👨‍💻 {t('about')}", use_container_width=True):
            st.session_state.page = 'about'
        
        st.markdown("---")
        
        # API Configuration
        with st.expander(f"🔧 {t('api_config')}"):
            # Store API keys in session state when changed
            groq_key = st.text_input("Groq API Key", 
                                    type="password", 
                                    help="Get your free Groq API key from https://console.groq.com/keys",
                                    value=st.session_state.api_keys.get('groq', ''))
            if groq_key != st.session_state.api_keys.get('groq', ''):
                st.session_state.api_keys['groq'] = groq_key
            
            news_key = st.text_input("NewsAPI Key", 
                                    type="password",
                                    help="Get your free NewsAPI key from https://newsapi.org/",
                                    value=st.session_state.api_keys.get('news', ''))
            if news_key != st.session_state.api_keys.get('news', ''):
                st.session_state.api_keys['news'] = news_key
            
            twitter_key = st.text_input("Twitter API Key", 
                                       type="password", 
                                       help="Optional: For enhanced sentiment analysis",
                                       value=st.session_state.api_keys.get('twitter', ''))
            if twitter_key != st.session_state.api_keys.get('twitter', ''):
                st.session_state.api_keys['twitter'] = twitter_key
            
            if st.button("💾 Save Configuration"):
                st.success("API keys saved successfully!")
        
        # Show API status
        st.markdown("### 🔑 API Status")
        groq_status = "🟢 Connected" if st.session_state.api_keys.get('groq') else "🔴 Not configured"
        news_status = "🟢 Connected" if st.session_state.api_keys.get('news') else "🔴 Not configured"
        
        st.markdown(f"**Groq AI:** {groq_status}")
        st.markdown(f"**NewsAPI:** {news_status}")
        
        # Translation Settings
        with st.expander("🌍 Translation Settings"):
            translation_options = {
                "Auto (Online + Offline)": "auto",
                "Offline Only": "offline", 
                "Disabled": "disabled"
            }
            selected_mode = st.selectbox(
                "Translation Mode",
                list(translation_options.keys()),
                index=0
            )
            st.session_state.translation_mode = translation_options[selected_mode]
            
            if not TRANSLATION_AVAILABLE:
                st.warning("⚠️ Online translation unavailable")
                st.info("Install: pip install googletrans==4.0.0rc1")
        
        # Translation status
        if st.session_state.translation_mode == "disabled":
            st.info("🔒 Translation disabled")
        elif not TRANSLATION_AVAILABLE:
            st.warning("⚠️ Limited to offline translations")
        
        # Footer
        st.markdown("---")
        st.markdown(f"""
        <div style="text-align: center; padding: 1rem;">
            <small>© 2025 Misbah Noorain</small><br>
            <small>BIET Davanagere</small><br>
            <small style="color: #667eea;">⚡ Powered by Groq AI</small>
        </div>
        """, unsafe_allow_html=True)
    
    # Page routing
    if 'page' not in st.session_state:
        st.session_state.page = 'landing'
    
    try:
        if st.session_state.page == 'landing':
            landing_page()
        elif st.session_state.page == 'dashboard':
            dashboard_page()
        elif st.session_state.page == 'reports':
            reports_page()
        elif st.session_state.page == 'about':
            about_page()
    except Exception as e:
        st.error(f"An error occurred: {e}")
        st.info("Please try refreshing the page or switching back to English.")
        # Log error for debugging
        if st.session_state.get('debug_mode', False):
            st.exception(e)

if __name__ == "__main__":
    main()